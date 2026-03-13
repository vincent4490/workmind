# -*- coding: utf-8 -*-
"""
需求智能体：导出为 Word（.docx）/ XMind（.xmind）
- 产品、开发角色任务 → Word
- 测试角色任务 → XMind
"""
import io
import re
import logging
from .requirement_xmind import build_to_bytes as xmind_build_to_bytes

logger = logging.getLogger(__name__)

# 任务类型 -> 中文名（用于文件名）
TASK_TYPE_LABELS = {
    'competitive_analysis': '竞品分析',
    'prd_draft': 'PRD撰写',
    'prd_refine': '需求完善',
    'requirement_analysis': '需求分析',
    'tech_design': '技术方案',
    'test_requirement_analysis': '测试需求分析',
}


def _safe_filename(title: str, task_type: str, ext: str) -> str:
    """生成安全文件名：标题前30字_任务类型.扩展名"""
    label = TASK_TYPE_LABELS.get(task_type, task_type)
    clean = re.sub(r'[<>:"/\\|?*]', '_', (title or '')[:30].strip()) or '需求'
    return f"{clean}_{label}.{ext}"


def export_docx(task) -> tuple[bytes, str]:
    """
    将任务导出为 Word (.docx)，适用于产品/开发角色任务。
    :return: (文件二进制, 建议文件名)
    """
    from docx import Document

    content = (task.result_md or getattr(task, 'raw_content', '') or '').strip()
    if not content:
        raise ValueError("该任务无文本内容可导出为 Word")

    doc = Document()
    title = (task.requirement_input or '')[:80].strip() or '需求智能体输出'
    task_label = TASK_TYPE_LABELS.get(task.task_type, task.task_type)
    doc.add_heading(f"{title}\n[{task_label}]", level=0)

    for block in content.split('\n\n'):
        block = block.strip()
        if not block:
            continue
        for line in block.split('\n'):
            line = line.strip()
            if not line:
                continue
            if line.startswith('### '):
                doc.add_heading(line[4:].strip(), level=3)
            elif line.startswith('## '):
                doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith('# '):
                doc.add_heading(line[2:].strip(), level=1)
            else:
                doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    filename = _safe_filename(task.requirement_input, task.task_type, 'docx')
    return buf.getvalue(), filename


def export_xmind(task) -> tuple[bytes, str]:
    """
    将任务 result_json 导出为 XMind，适用于测试角色任务。
    :return: (文件二进制, 建议文件名)
    """
    if not getattr(task, "result_json", None):
        raise ValueError("该任务无结构化结果可导出为 XMind")
    data = xmind_build_to_bytes(task)
    filename = _safe_filename(task.requirement_input, task.task_type, 'xmind')
    return data, filename
