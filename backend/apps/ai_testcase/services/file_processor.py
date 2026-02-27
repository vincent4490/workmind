# -*- coding: utf-8 -*-
"""
文件处理器：从上传文件中提取文字和图片
支持 .docx / .pdf / .txt / .md / .png / .jpg / .jpeg / .webp
"""
import base64
import io
import os
import re
import logging
from pathlib import Path

from PIL import Image

logger = logging.getLogger(__name__)

# ============ 约束常量 ============
MAX_FILES = 10                          # 最多上传文件数
MAX_FILE_SIZE = 10 * 1024 * 1024        # 单文件最大 10MB
MAX_TOTAL_TEXT_CHARS = 30000            # 提取文字总上限（字符），约15K tokens
MAX_IMAGES = 20                         # 图片总上限（含文档内嵌入的），约20-30K tokens
IMAGE_MAX_DIMENSION = 1024              # 图片最大边长（像素）
IMAGE_QUALITY = 80                      # JPEG 压缩质量

SUPPORTED_DOC_EXTS = {'.docx', '.pdf', '.txt', '.md'}
SUPPORTED_IMG_EXTS = {'.png', '.jpg', '.jpeg', '.webp'}
SUPPORTED_EXTS = SUPPORTED_DOC_EXTS | SUPPORTED_IMG_EXTS


class FileProcessError(Exception):
    """文件处理异常"""
    pass


class FileProcessor:
    """
    统一文件处理器

    process_files(files) 返回:
    {
        "texts": [{"source": "文件名", "content": "提取的文字"}],
        "images": [{"source": "文件名", "data": "base64", "mime": "image/png"}],
        "warnings": ["某某文件处理失败的提示"]
    }
    """

    def process_files(self, files) -> dict:
        """处理所有上传文件"""
        if len(files) > MAX_FILES:
            raise FileProcessError(f'最多上传 {MAX_FILES} 个文件，当前 {len(files)} 个')

        texts = []
        images = []
        warnings = []

        for f in files:
            filename = f.name if hasattr(f, 'name') else str(f)
            ext = Path(filename).suffix.lower()

            if ext not in SUPPORTED_EXTS:
                warnings.append(f'不支持的文件格式: {filename}')
                continue

            # 检查文件大小
            file_size = f.size if hasattr(f, 'size') else 0
            if file_size > MAX_FILE_SIZE:
                warnings.append(f'文件过大（>{MAX_FILE_SIZE // 1024 // 1024}MB）: {filename}')
                continue

            try:
                if ext == '.docx':
                    t, imgs = self._process_docx(f, filename)
                elif ext == '.pdf':
                    t, imgs = self._process_pdf(f, filename)
                elif ext == '.txt':
                    t, imgs = self._process_txt(f, filename)
                elif ext == '.md':
                    t, imgs = self._process_md(f, filename)
                elif ext in SUPPORTED_IMG_EXTS:
                    t, imgs = self._process_image(f, filename)
                else:
                    continue

                if t:
                    texts.append({'source': filename, 'content': t})
                images.extend(imgs)

            except Exception as e:
                logger.warning(f'[FileProcessor] 处理文件 {filename} 失败: {e}')
                warnings.append(f'文件处理失败: {filename} ({str(e)[:60]})')

        # 文字截断
        texts = self._truncate_texts(texts, MAX_TOTAL_TEXT_CHARS)

        # 图片数量限制
        if len(images) > MAX_IMAGES:
            warnings.append(f'图片数量超限，仅保留前 {MAX_IMAGES} 张（共提取 {len(images)} 张）')
            images = images[:MAX_IMAGES]

        return {'texts': texts, 'images': images, 'warnings': warnings}

    def process_local_files(self, file_paths: list) -> dict:
        """
        从磁盘路径读取文件并提取文本和图片（用于模块重新生成场景）

        Args:
            file_paths: 文件绝对路径列表，如 ["D:/workmind/.../uploads/1/xxx.pdf"]

        Returns:
            与 process_files 完全相同的格式:
            {"texts": [...], "images": [...], "warnings": [...]}
        """
        texts = []
        images = []
        warnings = []

        for filepath in file_paths:
            if not os.path.exists(filepath):
                warnings.append(f'源文件不存在（可能已被删除）: {os.path.basename(filepath)}')
                continue

            filename = os.path.basename(filepath)
            ext = Path(filename).suffix.lower()

            if ext not in SUPPORTED_EXTS:
                warnings.append(f'不支持的文件格式: {filename}')
                continue

            try:
                with open(filepath, 'rb') as f:
                    if ext == '.docx':
                        t, imgs = self._process_docx(f, filename)
                    elif ext == '.pdf':
                        t, imgs = self._process_pdf(f, filename)
                    elif ext == '.txt':
                        t, imgs = self._process_txt(f, filename)
                    elif ext == '.md':
                        t, imgs = self._process_md(f, filename)
                    elif ext in SUPPORTED_IMG_EXTS:
                        t, imgs = self._process_image(f, filename)
                    else:
                        continue

                    if t:
                        texts.append({'source': filename, 'content': t})
                    images.extend(imgs)

            except Exception as e:
                logger.warning(f'[FileProcessor] 处理本地文件 {filepath} 失败: {e}')
                warnings.append(f'文件处理失败: {filename} ({str(e)[:60]})')

        # 文字截断
        texts = self._truncate_texts(texts, MAX_TOTAL_TEXT_CHARS)

        # 图片数量限制
        if len(images) > MAX_IMAGES:
            warnings.append(f'图片数量超限，仅保留前 {MAX_IMAGES} 张（共提取 {len(images)} 张）')
            images = images[:MAX_IMAGES]

        return {'texts': texts, 'images': images, 'warnings': warnings}

    # ======================== 各类型处理 ========================

    def _process_docx(self, file, filename: str) -> tuple:
        """提取 docx 的文字和嵌入图片"""
        import docx

        file_bytes = self._read_file_bytes(file)
        doc = docx.Document(io.BytesIO(file_bytes))

        # 提取文字
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text = '\n'.join(paragraphs)

        # 提取表格中的文字
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text += '\n' + row_text

        # 提取嵌入图片
        images = []
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    img_data = rel.target_part.blob
                    img_b64 = self._compress_and_encode_image(img_data, f'{filename}_embedded')
                    if img_b64:
                        images.append(img_b64)
                except Exception as e:
                    logger.debug(f'[FileProcessor] docx 图片提取失败: {e}')

        return text, images

    def _process_pdf(self, file, filename: str) -> tuple:
        """提取 pdf 的文字和嵌入图片"""
        import pdfplumber
        import fitz  # PyMuPDF

        file_bytes = self._read_file_bytes(file)

        # 用 pdfplumber 提取文字
        text_parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text() or ''
                if page_text.strip():
                    text_parts.append(page_text)

        text = '\n'.join(text_parts)

        # 用 PyMuPDF 提取嵌入图片
        images = []
        pdf_doc = fitz.open(stream=file_bytes, filetype='pdf')
        for page_idx in range(len(pdf_doc)):
            page = pdf_doc[page_idx]
            image_list = page.get_images(full=True)

            for img_info in image_list:
                xref = img_info[0]
                try:
                    base_image = pdf_doc.extract_image(xref)
                    img_data = base_image['image']
                    # 跳过太小的图片（可能是 logo/图标）
                    if len(img_data) < 2000:
                        continue
                    img_b64 = self._compress_and_encode_image(
                        img_data, f'{filename}_page{page_idx + 1}'
                    )
                    if img_b64:
                        images.append(img_b64)
                except Exception as e:
                    logger.debug(f'[FileProcessor] PDF 图片提取失败 (page {page_idx}): {e}')

            # 如果某页文字很少但有大量内容（可能是纯图片页），渲染该页为图片
            page_text = text_parts[page_idx] if page_idx < len(text_parts) else ''
            if len(page_text.strip()) < 50 and not image_list:
                try:
                    pix = page.get_pixmap(dpi=150)
                    img_data = pix.tobytes('png')
                    if len(img_data) > 5000:  # 忽略空白页
                        img_b64 = self._compress_and_encode_image(
                            img_data, f'{filename}_page{page_idx + 1}_render'
                        )
                        if img_b64:
                            images.append(img_b64)
                except Exception as e:
                    logger.debug(f'[FileProcessor] PDF 页面渲染失败 (page {page_idx}): {e}')

        pdf_doc.close()
        return text, images

    def _process_txt(self, file, filename: str) -> tuple:
        """读取纯文本"""
        file_bytes = self._read_file_bytes(file)
        # 尝试多种编码
        for encoding in ('utf-8', 'gbk', 'gb2312', 'latin-1'):
            try:
                text = file_bytes.decode(encoding)
                return text, []
            except (UnicodeDecodeError, LookupError):
                continue
        return file_bytes.decode('utf-8', errors='replace'), []

    def _process_md(self, file, filename: str) -> tuple:
        """读取 Markdown 文字 + 解析图片引用"""
        file_bytes = self._read_file_bytes(file)
        text = file_bytes.decode('utf-8', errors='replace')

        # 解析图片引用
        images = []
        # 匹配 ![alt](url) 和 <img src="url">
        img_urls = re.findall(r'!\[.*?\]\((https?://[^)]+)\)', text)
        img_urls += re.findall(r'<img[^>]+src=["\']?(https?://[^\s"\'>]+)', text)

        for url in img_urls[:5]:  # 最多下载 5 张远程图片
            try:
                img_data = self._download_image(url)
                if img_data:
                    img_b64 = self._compress_and_encode_image(img_data, url.split('/')[-1])
                    if img_b64:
                        images.append(img_b64)
            except Exception as e:
                logger.debug(f'[FileProcessor] 下载 md 图片失败 {url}: {e}')

        return text, images

    def _process_image(self, file, filename: str) -> tuple:
        """处理独立图片文件"""
        file_bytes = self._read_file_bytes(file)
        img_b64 = self._compress_and_encode_image(file_bytes, filename)
        if img_b64:
            return '', [img_b64]
        return '', []

    # ======================== 工具方法 ========================

    def _read_file_bytes(self, file) -> bytes:
        """读取上传文件的字节内容"""
        if hasattr(file, 'read'):
            file.seek(0)
            data = file.read()
            file.seek(0)
            return data
        return bytes(file)

    def _compress_and_encode_image(self, img_data: bytes, source: str) -> dict | None:
        """
        压缩图片并编码为 base64
        返回: {"source": "...", "data": "base64...", "mime": "image/jpeg"}
        """
        try:
            img = Image.open(io.BytesIO(img_data))

            # 转换 RGBA/P 模式为 RGB（JPEG 不支持透明通道）
            if img.mode in ('RGBA', 'P', 'LA'):
                background = Image.new('RGB', img.size, (255, 255, 255))
                if img.mode == 'P':
                    img = img.convert('RGBA')
                background.paste(img, mask=img.split()[-1] if 'A' in img.mode else None)
                img = background
            elif img.mode != 'RGB':
                img = img.convert('RGB')

            # 缩放
            w, h = img.size
            if max(w, h) > IMAGE_MAX_DIMENSION:
                ratio = IMAGE_MAX_DIMENSION / max(w, h)
                new_size = (int(w * ratio), int(h * ratio))
                img = img.resize(new_size, Image.LANCZOS)

            # 编码为 JPEG base64
            buffer = io.BytesIO()
            img.save(buffer, format='JPEG', quality=IMAGE_QUALITY)
            b64_str = base64.b64encode(buffer.getvalue()).decode('utf-8')

            return {
                'source': source,
                'data': b64_str,
                'mime': 'image/jpeg'
            }

        except Exception as e:
            logger.warning(f'[FileProcessor] 图片压缩编码失败 ({source}): {e}')
            return None

    def _download_image(self, url: str, timeout: int = 10) -> bytes | None:
        """下载远程图片"""
        import urllib.request
        try:
            req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
            with urllib.request.urlopen(req, timeout=timeout) as resp:
                data = resp.read()
                if len(data) > MAX_FILE_SIZE:
                    return None
                return data
        except Exception as e:
            logger.debug(f'[FileProcessor] 下载图片失败 {url}: {e}')
            return None

    def _truncate_texts(self, texts: list, max_chars: int) -> list:
        """文字总量超限时按比例截断"""
        total = sum(len(t['content']) for t in texts)
        if total <= max_chars:
            return texts

        # 按比例分配每个文件的字符配额
        result = []
        for t in texts:
            ratio = len(t['content']) / total
            allowed = int(max_chars * ratio)
            content = t['content']
            if len(content) > allowed:
                content = content[:allowed] + f'\n\n... [文件 {t["source"]} 内容过长，已截断] ...'
            result.append({'source': t['source'], 'content': content})
        return result
